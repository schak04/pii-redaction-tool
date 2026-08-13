from pathlib import Path
from typing import List, Union
from docx import Document

def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def extract_paragraphs(docx_path: Union[str, Path]) -> List[str]:
    doc = Document(str(docx_path))
    return [p.text for p in iter_all_paragraphs(doc)]

def extract_text(docx_path: Union[str, Path]) -> str:
    paragraphs = extract_paragraphs(docx_path)
    return "\n".join(paragraphs)
