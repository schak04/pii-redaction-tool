import tempfile
import unittest
from pathlib import Path
from docx import Document
from fastapi.testclient import TestClient
from src.api import app

class TestAPI(unittest.TestCase):
    def setUp(self) -> None:
        self.client = TestClient(app)
        self.temp_dir = tempfile.TemporaryDirectory()
        self.docx_path = Path(self.temp_dir.name) / "test_api.docx"

        doc = Document()
        doc.add_paragraph("Contact Knight at knight@hallow-nest.org.")
        doc.save(str(self.docx_path))

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def test_health_endpoint(self) -> None:
        response = self.client.get("/health")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["status"], "ok")

    def test_redact_endpoint(self) -> None:
        with open(self.docx_path, "rb") as f:
            files = {"file": ("test_api.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            response = self.client.post("/redact?threshold=0.3", files=files)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        res_path = Path(self.temp_dir.name) / "res.docx"
        with open(res_path, "wb") as f:
            f.write(response.content)

        res_doc = Document(str(res_path))
        self.assertNotIn("knight@hallow-nest.org", res_doc.paragraphs[0].text)
        self.assertIn("<EMAIL_ADDRESS>", res_doc.paragraphs[0].text)

if __name__ == "__main__":
    unittest.main()
