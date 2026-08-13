import tempfile
import unittest
from pathlib import Path
from docx import Document
from presidio_analyzer import RecognizerResult
from src.redactor import DOCXRedactor, redact_paragraph_runs

class TestDOCXRedactor(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.doc_path = Path(self.temp_dir.name) / "sample.docx"
        self.out_path = Path(self.temp_dir.name) / "redacted.docx"

        doc = Document()
        p1 = doc.add_paragraph()
        run1 = p1.add_run("Contact Papyrus at ")
        run2 = p1.add_run("papyrus@undertale.com")
        run2.bold = True
        p1.add_run(" today.")

        doc.add_paragraph("Sans lives in Snowdin.")
        doc.save(str(self.doc_path))

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def test_single_run_redaction(self) -> None:
        redactor = DOCXRedactor(self.doc_path)
        p1_detections = [
            RecognizerResult(entity_type="EMAIL_ADDRESS", start=19, end=40, score=1.0)
        ]
        p2_detections = [
            RecognizerResult(entity_type="PERSON", start=0, end=4, score=0.9)
        ]

        redactor.redact([p1_detections, p2_detections])
        redactor.save(self.out_path)

        redacted_doc = Document(str(self.out_path))
        text1 = redacted_doc.paragraphs[0].text
        text2 = redacted_doc.paragraphs[1].text

        self.assertIn("<EMAIL_ADDRESS>", text1)
        self.assertNotIn("papyrus@undertale.com", text1)
        self.assertIn("<PERSON>", text2)
        self.assertNotIn("Sans", text2)

    def test_multi_run_spanning_redaction(self) -> None:
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("Hero ")
        r2 = p.add_run("Knight")
        
        detections = [
            RecognizerResult(entity_type="PERSON", start=0, end=11, score=1.0)
        ]
        redact_paragraph_runs(p, detections)
        self.assertEqual(p.text, "<PERSON>")

if __name__ == "__main__":
    unittest.main()
