import tempfile
import unittest
from pathlib import Path
from docx import Document
from src.extractor import extract_paragraphs, extract_text

class TestExtractor(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.file_path = Path(self.temp_dir.name) / "test_prospectus.docx"
        
        doc = Document()
        doc.add_paragraph("Hornet resides in Hallownest.")
        doc.add_paragraph("Contact Knight at knight@hallow-nest.org or call +1 555 0199.")
        
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].paragraphs[0].text = "Officer Name: Solaire"
        table.rows[0].cells[1].paragraphs[0].text = "Email: solaire@astora.org"

        doc.save(str(self.file_path))

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def test_extract_paragraphs(self) -> None:
        paragraphs = extract_paragraphs(self.file_path)
        self.assertEqual(len(paragraphs), 4)
        self.assertEqual(paragraphs[0], "Hornet resides in Hallownest.")
        self.assertIn("Officer Name: Solaire", paragraphs)

    def test_extract_text(self) -> None:
        text = extract_text(self.file_path)
        self.assertIn("Hornet resides in Hallownest.", text)
        self.assertIn("solaire@astora.org", text)

if __name__ == "__main__":
    unittest.main()
