import sys
import subprocess
import tempfile
import unittest
from pathlib import Path
from docx import Document
from src.pipeline import PIIRedactionPipeline

class TestPipeline(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.in_doc = Path(self.temp_dir.name) / "prospectus.docx"
        self.out_doc = Path(self.temp_dir.name) / "sanitized.docx"

        doc = Document()
        doc.add_paragraph("Solaire of Astora contacted solaire@example.com.")
        doc.add_paragraph("The server IP is 192.168.1.1.")
        doc.save(str(self.in_doc))

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def test_end_to_end_pipeline(self) -> None:
        pipeline = PIIRedactionPipeline(score_threshold=0.3)
        pipeline.process_document(self.in_doc, self.out_doc)

        self.assertTrue(self.out_doc.exists())

        res_doc = Document(str(self.out_doc))
        text1 = res_doc.paragraphs[0].text
        text2 = res_doc.paragraphs[1].text

        self.assertNotIn("solaire@example.com", text1)
        self.assertIn("<EMAIL_ADDRESS>", text1)
        self.assertNotIn("192.168.1.1", text2)
        self.assertIn("<IP_ADDRESS>", text2)

    def test_cli_execution(self) -> None:
        cmd = [
            sys.executable,
            "src/redact.py",
            str(self.in_doc),
            "-o",
            str(self.out_doc),
            "-t",
            "0.3",
        ]
        res = subprocess.run(cmd, capture_output=True, text=True)
        self.assertEqual(res.returncode, 0, msg=f"CLI failed with error:\n{res.stderr}")
        self.assertTrue(self.out_doc.exists())

if __name__ == "__main__":
    unittest.main()
